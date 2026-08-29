import json

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

MAX_STRUCTURAL_UNITS = 100_000
MAX_PARAGRAPHS = 20_000
MAX_TABLES = 2_000
MAX_ROWS_PER_TABLE = 5_000
MAX_CELLS_PER_ROW = 256
MAX_SOURCES = 20_000
MAX_SOURCE_CODE_POINTS = 100_000
MAX_DOCUMENT_CODE_POINTS = 2_000_000


def add_source(sources, source, state):
    content = source["content"]
    if len(content) > MAX_SOURCE_CODE_POINTS:
        raise ValueError("source_limit")
    state["points"] += len(content)
    if state["points"] > MAX_DOCUMENT_CODE_POINTS:
        raise ValueError("document_limit")
    if len(sources) >= MAX_SOURCES:
        raise ValueError("source_count_limit")
    sources.append(source)


def count_unit(state):
    state["units"] += 1
    if state["units"] > MAX_STRUCTURAL_UNITS:
        raise ValueError("structure_limit")


def add_table(table, table_index, sources, state):
    for row_index, row in enumerate(table.rows, start=1):
        count_unit(state)
        if row_index > MAX_ROWS_PER_TABLE:
            raise ValueError("row_limit")
        for column_index, cell in enumerate(row.cells, start=1):
            count_unit(state)
            if column_index > MAX_CELLS_PER_ROW:
                raise ValueError("column_limit")
            content = cell.text.strip()
            if content:
                source = {"kind": "table_cell", "table": table_index, "row": row_index,
                          "column": column_index, "content": content}
                add_source(sources, source, state)


def add_body(document, sources, state):
    paragraph_index = 0
    table_index = 0
    for item in document.iter_inner_content():
        count_unit(state)
        if isinstance(item, Paragraph):
            paragraph_index += 1
            if paragraph_index > MAX_PARAGRAPHS:
                raise ValueError("paragraph_limit")
            content = item.text.strip()
            if content:
                add_source(sources, {"kind": "paragraph", "paragraph": paragraph_index,
                                     "content": content}, state)
        elif isinstance(item, Table):
            table_index += 1
            if table_index > MAX_TABLES:
                raise ValueError("table_limit")
            add_table(item, table_index, sources, state)


def parse_docx(path):
    document = Document(path)
    sources: list[dict[str, object]] = []
    state = {"units": 0, "points": 0}
    add_body(document, sources, state)
    if not sources:
        raise ValueError("no_text")
    return json.dumps(
        {"schema_version": "1", "format": "docx", "sources": sources},
        ensure_ascii=False,
        separators=(",", ":"),
    )


parse_docx("/tmp/aethelgard-source.docx")
